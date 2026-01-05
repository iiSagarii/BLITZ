import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import os
import docx
import re
import subprocess
import threading
import tkinter.ttk as ttk
import sys # Added for console redirection
import io  # Added for console redirection

# Hardcoded paths for base and output files (for TOE type processing)
BASE_TSS_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "sys_inst/base_TSS.txt")
OUTPUT_TSS_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)),"sys_inst/System_inst-TSS.txt")

# Hardcoded paths for SD documents (scalable structure)
SD_OPTIONS = {
    "Protection_Profiles": {
        "NDcPP_v3.0": os.path.join(os.path.dirname(os.path.abspath(__file__)), "DB/NDcPP_v3.0.docx"),
        # "NDcPP_v2.2e": os.path.join(os.path.dirname(os.path.abspath(__file__)), "DB/NDcPP_v2.2e.docx")
    },
    "Packages": {
        "coming soon ..."
        # "PKG_SSH_v1.0": os.path.join(os.path.dirname(os.path.abspath(__file__)), "DB/PKG_SSH_v1.0.docx")
    },
    "Modules": {
        "coming soon ..."
    }
}

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "ephemeral")

### Requirements Processor Class
class RequirementsProcessor:
    def __init__(self):
        self.st_data = {}
        self.sfr_data = {}
        self.sd_data = {}

    def extract_st_data(self, doc_path):
        """Extract requirements and SFR content from ST document"""
        try:
            doc = docx.Document(doc_path)
            requirements_data = {}  # TSS data from table or section
            sfr_data = {}          # SFR data from SFR section

            # Helper function to iterate through paragraphs and tables in order
            from docx.oxml import CT_P, CT_Tbl
            def iter_block_items(parent):
                for child in parent.element.body:
                    if isinstance(child, CT_P):
                        yield docx.text.paragraph.Paragraph(child, parent)
                    elif isinstance(child, CT_Tbl):
                        yield docx.table.Table(child, parent)

            # Helper function to extract full text from a cell, including nested tables
            def extract_full_text_from_cell(cell):
                text = ""
                for para in cell.paragraphs:
                    text += para.text + "\n"
                for table in cell.tables:
                    text += "This is the data from a table within the TSS text\n"
                    for row in table.rows:
                        for cell in row.cells:
                            text += extract_full_text_from_cell(cell) + "\n"
                return text.strip()

            # Extract SFR content from "Security Functional Requirements" section
            in_sfr_section = False
            current_req = None
            current_content = []
            section_style = None
            for para in doc.paragraphs:
                if para.style.name.startswith('Heading 2'):
                    if re.search(r'security\s+functional\s+requirements', para.text.strip(), re.IGNORECASE):
                        in_sfr_section = True
                        section_style = para.style.name
                        continue
                    elif in_sfr_section and para.style.name == section_style:
                        if current_req and current_content:
                            sfr_data[current_req] = "\n".join(current_content)
                        in_sfr_section = False
                        current_req = None
                        current_content = []
                        continue

                if in_sfr_section:
                    if para.style.name == 'Heading 4':  # SFR sub-sections are 'Heading 4'
                        match = re.match(r'([A-Z]{3}_[A-Z0-9._]+(?:/[A-Za-z]+)?)\s+.*', para.text.strip())
                        if match:
                            if current_req and current_content:
                                sfr_data[current_req] = "\n".join(current_content)
                            current_req = match.group(1)
                            current_content = []
                        elif current_req: # If it's Heading 4 but doesn't match SFR pattern, append to previous
                            current_content.append(para.text.strip())
                    elif current_req:
                        current_content.append(para.text.strip())

            if in_sfr_section and current_req and current_content:
                sfr_data[current_req] = "\n".join(current_content)

            # Step 1: Try to extract TSS from the table "TOE Summary Specification SFR Description"
            tss_table_found = False
            for table in doc.tables:
                if len(table.columns) >= 2:
                    try:
                        header_row = table.rows[0]
                        if len(header_row.cells) >= 2 and 'tss description' in header_row.cells[1].text.strip().lower():
                            tss_table_found = True
                            for row in table.rows[1:]:  # Skip header row
                                if len(row.cells) >= 2:
                                    requirement_cell = row.cells[0].text.strip()
                                    requirements = re.findall(r'[A-Z]{3}_[A-Z0-9._]+(?:/[A-Za-z]+)?', requirement_cell)
                                    if requirements:
                                        tss_desc = extract_full_text_from_cell(row.cells[1])
                                        for req in requirements:
                                            requirements_data[req] = tss_desc
                            break # Stop after processing the first matching table
                    except IndexError:
                        # Handle potential index errors if table structure is unexpected
                        print(f"Warning: Skipping a table due to unexpected structure.")
                        continue

            # Step 2: If no table is found, extract TSS from "TOE Summary Specifications" section
            if not tss_table_found:
                in_tss_section = False
                current_subsection = None
                subsection_content = []
                main_tss_heading_style = None # To detect end of TSS section

                for block in iter_block_items(doc):
                    if isinstance(block, docx.text.paragraph.Paragraph):
                        para = block
                        text = para.text.strip()
                        if not text:
                            continue

                        # Detect start of TSS section
                        if para.style.name.startswith('Heading') and re.search(r'toe\s+summary\s+specifications', text, re.IGNORECASE):
                            in_tss_section = True
                            main_tss_heading_style = para.style.name
                            current_subsection = None # Reset subsection when main heading found
                            subsection_content = []
                            continue

                        if in_tss_section:
                            # Detect end of TSS section (e.g., start of SFR section or another main heading)
                            if para.style.name == main_tss_heading_style and not re.search(r'toe\s+summary\s+specifications', text, re.IGNORECASE):
                                if current_subsection and subsection_content:
                                     requirements_data[current_subsection] = "\n".join(subsection_content).strip()
                                in_tss_section = False
                                break # Exit loop once TSS section ends
                            elif para.style.name.startswith('Heading') and re.search(r'security\s+functional\s+requirements', text, re.IGNORECASE):
                                if current_subsection and subsection_content:
                                     requirements_data[current_subsection] = "\n".join(subsection_content).strip()
                                in_tss_section = False
                                break # Exit loop once TSS section ends

                            # Detect TSS subsections (assuming they are one level below main TSS heading)
                            # This logic might need adjustment based on actual document structure
                            is_subsection_heading = False
                            if para.style.name.startswith('Heading'):
                                match = re.match(r'([A-Z]{3}_[A-Z0-9._]+(?:/[A-Za-z]+)?)\s+.*', text)
                                if match:
                                    is_subsection_heading = True
                                    if current_subsection and subsection_content:
                                        requirements_data[current_subsection] = "\n".join(subsection_content).strip()
                                    current_subsection = match.group(1)
                                    subsection_content = [] # Start new content list for the subsection
                                    # Optionally add the heading itself to the content
                                    # subsection_content.append(text)
                                    continue # Skip adding heading text as normal content

                            if current_subsection and not is_subsection_heading:
                                subsection_content.append(text)

                    elif isinstance(block, docx.table.Table) and in_tss_section and current_subsection:
                        table_text = f"--- Table Data for {current_subsection} ---\n"
                        for row in block.rows:
                            for cell in row.cells:
                                table_text += extract_full_text_from_cell(cell).replace("\n", " | ") + "\t" # Simple table formatting
                            table_text += "\n"
                        table_text += f"--- End Table Data for {current_subsection} ---\n"
                        subsection_content.append(table_text)

                # Save the last subsection's content if the loop finished while still in TSS section
                if in_tss_section and current_subsection and subsection_content:
                    requirements_data[current_subsection] = "\n".join(subsection_content).strip()


            # Error control: Check if all SFRs have TSS entries
            missing_sfrs = set(sfr_data.keys()) - set(requirements_data.keys())
            if missing_sfrs:
                raise ValueError("Mismatch between the 'TSS' and 'Security Functional Requirements' section for: " + ", ".join(missing_sfrs) + ". Please correct the ST document.")

            # Check if TSS entries exist that are not in SFRs (optional warning)
            extra_tss = set(requirements_data.keys()) - set(sfr_data.keys())
            if extra_tss:
                 print(f"Warning: Found TSS entries without corresponding SFRs: {', '.join(extra_tss)}")


            return requirements_data, sfr_data
        except ValueError as e:
            # messagebox.showerror("Error", str(e)) # Handled in calling function
            raise
        except Exception as e:
            messagebox.showerror("Error", f"Error processing ST document: {str(e)}")
            return {}, {} # Return empty dicts on failure


    def extract_sd_data(self, sd_paths, st_requirements):
        """Extract content from SD documents by recognizing bold text as headings"""
        sd_data = {req: [] for req in st_requirements}
        for sd_path in sd_paths:
            try:
                doc = docx.Document(sd_path)
                current_section = None
                section_content = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        # Check if the paragraph is bold (heading)
                        if all(run.bold for run in para.runs if run.text.strip()):
                            # If we were in a section, save its content
                            if current_section and section_content:
                                sd_data[current_section].append("\n".join(section_content))
                            # Check if this bold text matches any ST requirement
                            for req in st_requirements:
                                if re.search(rf'\b{re.escape(req)}\b', para.text.strip()):
                                    current_section = req
                                    section_content = []
                                    break
                            else:
                                current_section = None  # Not a relevant heading
                        elif current_section:
                            # Collect content under the current section
                            section_content.append(para.text.strip())
                # Save the last section's content
                if current_section and section_content:
                    sd_data[current_section].append("\n".join(section_content))
            except Exception as e:
                messagebox.showerror("Error", f"Error processing SD document {sd_path}: {str(e)}")
        for req in sd_data:
            sd_data[req] = "\n\n".join(sd_data[req]) if sd_data[req] else "No description found"
        return sd_data


    def process_files(self, st_path, sd_paths):
        """Process ST and SD files, saving results in chunks"""
        try:
            self.st_data, self.sfr_data = self.extract_st_data(st_path)
        except ValueError as e:
            return False, str(e)
        if not self.st_data:
            return False, "No valid data found in ST document"
        self.sd_data = self.extract_sd_data(sd_paths, self.st_data.keys())
        if not self.sd_data:
            return False, "No valid data found in SD documents"
        missing_requirements = [req for req in self.st_data if self.sd_data[req] == "No description found"]
        if missing_requirements:
            error_message = "Incorrect PP/MODS/PKGs selected. Please re-select the correct PP/MODS/PKGs for this evaluation.\nMissing requirements: " + ", ".join(missing_requirements) + ".\n\nPlease verify your PP/MODS/PKGs selection."
            return False, error_message
        if not self.st_data:
            # This case might be hit if extract_st_data returns empty dicts due to non-ValueError exceptions handled within it
            return False, "No valid TSS/SFR data found in the uploaded document. Please verify the document and re-upload the correct ST document."

        print("ST data extracted successfully.")
        print(f"SFRs found: {len(self.sfr_data)}")
        print(f"TSS entries found: {len(self.st_data)}")


        self.sd_data = self.extract_sd_data(sd_paths, self.st_data.keys())
        print("SD data extraction finished.")

        #missing_requirements = [req for req in self.st_data if self.sd_data[req] == "No description found in selected SD documents"]
        #if missing_requirements:            
         #   error_message = "Warning: No description found in the selected PP/MOD/PKG for the following requirements: " + ", ".join(missing_requirements) + ".\n\nPlease verify your PP/MODS/PKGs selection or the content of the SD documents. Processing will continue, but output for these requirements will be incomplete."
            # raise ValueError(error_message)
            # Decide if this should be a hard stop or just a warning
            # For now, let's make it a warning and continue, but return the message
          #  print(error_message) # Also print warning to console output
           # return False, error_message # Comment/Uncomment this line to make it a hard stop
                       
        
        os.makedirs(OUTPUT_DIR, exist_ok=True)
        print(f"Output directory: {OUTPUT_DIR}")

        requirements = sorted(self.st_data.keys())
        chunk_size = 3
        num_chunks = (len(requirements) + chunk_size - 1) // chunk_size
        print(f"Processing {len(requirements)} requirements in {num_chunks} chunk(s)...")


        for i in range(0, len(requirements), chunk_size):
            chunk = requirements[i:i + chunk_size]
            chunk_num = i // chunk_size + 1
            output_file = os.path.join(OUTPUT_DIR, f"user_prompt_TSS-{chunk_num}.txt")
            print(f"Writing chunk {chunk_num} to {output_file}...")
            try:
                with open(output_file, "w", encoding="utf-8") as f:
                    for req in chunk:
                        sfr_statement = self.sfr_data.get(req, "No SFR content found in ST") # Get SFR from stored data
                        tss_text = self.st_data[req]
                        sd_desc = self.sd_data[req] # Already contains "No description found..." if applicable

                        #f.write(f"Requirement: {req}\n")
                        #f.write("="*len(f"Requirement: {req}") + "\n\n")

                        f.write(f"SFR statement for {req}:\n{sfr_statement}\n\n")
                        f.write("-" * 20 + "\n\n")
                        f.write(f"TSS text for {req}:\n{tss_text}\n\n")
                        f.write("-" * 20 + "\n\n")
                        f.write(f"{sd_desc}\n\n")
                        f.write("=" * 80 + "\n\n")
            except Exception as e:
                 error_msg = f"Error writing chunk {chunk_num} to file {output_file}: {e}"
                 print(error_msg)
                 # Decide if one chunk error should stop everything
                 return False, error_msg # Stop processing if a chunk fails to write

        print("Finished writing all chunks.")
        return True, "TSS/SFR/SD data processing completed successfully" # Return success True


### TOE Type File Processing Function
def process_file(base_path, output_path, selection):
    """
    Process a base file, replace the placeholder with the user's selection, and save to the output file.
    Args:
        base_path (str): Path to the base file
        output_path (str): Path to save the modified file
        selection (str): User's selection ("STANDALONE" or "DISTRIBUTED")
    """
    try:
        with open(base_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        placeholder = "## In this case the TOE is <STANDALONE/DISTRIBUTED>."
        new_line = f"## In this case the TOE is {selection}."
        modified_lines = [new_line + '\n' if line.strip() == placeholder.strip() else line for line in lines]
        with open(output_path, 'w', encoding='utf-8') as f:
            f.writelines(modified_lines)
        print(f"Generated file: {output_path} with TOE type: {selection}") # Added print statement
    except FileNotFoundError:
        messagebox.showerror("Error", f"Base file not found: {base_path}")
        print(f"Error: Base file not found at {base_path}")
        raise # Re-raise error to be caught by caller if necessary
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred during file generation: {str(e)}")
        print(f"Error generating file {output_path}: {e}")
        raise # Re-raise error


### Main Application Class
class BlitzApp(tk.Tk):
    def __init__(self):
        tk.Tk.__init__(self)
        self.title("BLITZ")
        self.geometry("1000x700")
        self.configure(bg="black")
        try:
            # Ensure the path to icon.png is correct relative to the script location
            icon_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "UI/icon.png")
            if os.path.exists(icon_path):
                self.iconphoto(True, tk.PhotoImage(file=icon_path))
            else:
                 print(f"Warning: Icon file not found at {icon_path}")
        except tk.TclError as e:
            print(f"Error setting icon (TclError): {e}. Ensure UI/icon.png exists and is a valid PNG.")
        except Exception as e:
            print(f"Error setting icon: {e}")


        # Create container for frames
        self.container = tk.Frame(self, bg="black")
        self.container.pack(side="top", fill="both", expand=True)
        self.container.grid_rowconfigure(0, weight=1)
        self.container.grid_columnconfigure(0, weight=1)

        # Initialize frames
        self.video_frame = VideoFrame(self.container, self)
        self.choice_frame = ChoiceFrame(self.container, self)
        # Pass controller to ProcessingFrame
        self.processing_frame = ProcessingFrame(self.container, self)
        self.set_key_frame = SetKeyFrame(self.container, self)


        # Place frames in the container
        self.video_frame.grid(row=0, column=0, sticky="nsew")
        self.choice_frame.grid(row=0, column=0, sticky="nsew")
        self.processing_frame.grid(row=0, column=0, sticky="nsew")
        self.set_key_frame.grid(row=0, column=0, sticky="nsew")

        # Show the video frame initially
        self.show_frame(self.video_frame)

    def show_frame(self, frame):
        """Raise the specified frame to the top"""
        frame.tkraise()

### Video Frame
class VideoFrame(tk.Frame):
    def __init__(self, parent, controller):
        tk.Frame.__init__(self, parent, bg="black")
        self.video_label = tk.Label(self, bg="black")
        self.video_label.pack(pady=20)
        self.player = None # Initialize player attribute
        try:
            import tkvideo
            # Ensure the path to video.mp4 is correct relative to the script location
            video_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "UI/video.mp4")
            if os.path.exists(video_path):
                self.player = tkvideo.tkvideo(video_path, self.video_label, loop=1, size=(800, 450))
                self.player.play()
            else:
                print(f"Warning: Video file not found at {video_path}")
                self.video_label.config(text="Video not found.", fg="red") # Display error message
        except ImportError:
            print("tkvideo not installed. Please install it using 'pip install tkvideo'")
            self.video_label.config(text="tkvideo library not found. Please install.", fg="red")
        except Exception as e:
            print(f"Error playing video: {e}")
            self.video_label.config(text=f"Error playing video: {e}", fg="red")


        continue_frame = tk.Frame(self, bg="cyan2")
        continue_frame.pack(pady=10)
        continue_button = tk.Button(continue_frame, text="Continue", command=lambda: controller.show_frame(controller.choice_frame), fg="cyan2", bg="black", relief="flat", font=("Arial", 10))
        continue_button.pack(padx=1, pady=1)

    # Optional: Add method to stop video when switching frames if tkvideo supports it
    def stop_video(self):
         if self.player:
             try:
                 self.player.stop() # Or pause(), depending on tkvideo API
             except Exception as e:
                 print(f"Could not stop video: {e}")


### Choice Frame (Window 1)
class ChoiceFrame(tk.Frame):
    def __init__(self, parent, controller):
        tk.Frame.__init__(self, parent, bg="gray8")
        label = tk.Label(self, text="Please choose an option:", fg="white", bg="gray8", font=("Arial", 13, "bold"))
        label.pack(pady=20)
        # Set-key button
        set_key_frame = tk.Frame(self, bg="cyan2")
        set_key_frame.pack(pady=10)
        set_key_button = tk.Button(set_key_frame, font=("Arial", 10), text="Set API Key", command=lambda: controller.show_frame(controller.set_key_frame), fg="cyan2", bg="black", relief="flat", width=20)
        set_key_button.pack(padx=1, pady=1)
        # Validate AAs for ST-TSS button
        validate_frame = tk.Frame(self, bg="cyan2")
        validate_frame.pack(pady=10)
        validate_button = tk.Button(validate_frame, font=("Arial", 10), text="Validate TSS", command=lambda: controller.show_frame(controller.processing_frame), fg="cyan2", bg="black", relief="flat", width=20)
        validate_button.pack(padx=1, pady=1)

### Set Key Frame (Window 3)
class SetKeyFrame(tk.Frame):
    def __init__(self, parent, controller):
        tk.Frame.__init__(self, parent, bg="gray8")
        self.controller = controller
        # Back button
        back_frame = tk.Frame(self, bg="dark slate gray")
        back_frame.pack(anchor="nw", padx=10, pady=10)
        self.back_button = tk.Button(back_frame, text="Back", command=lambda: controller.show_frame(controller.choice_frame), fg="dark slate gray", bg="black", relief="flat") # Store reference
        self.back_button.pack(padx=1, pady=1)
        # Key input
        label = tk.Label(self, text="Please provide your API key:", fg="white", bg="gray8", font=("Arial", 13))
        label.pack(pady=20)

        key_entry_frame = tk.Frame(self, bg="gray8") # Frame to hold entry and placeholder logic
        key_entry_frame.pack(pady=10)

        self.key_entry = tk.Entry(key_entry_frame, fg="dark grey", bg="black", font=("Arial", 10), width=50) # Adjust width as needed
        self.key_entry.insert(0, "Paste your key here.") # Use index 0 for insert
        self.key_entry.bind("<FocusIn>", self.on_entry_click)
        self.key_entry.bind("<FocusOut>", self.on_focus_out)
        self.key_entry.pack() # Pack entry inside its frame

        # OK button
        self.ok_button_frame = tk.Frame(self, bg="cyan2")
        self.ok_button_frame.pack(pady=10)
        self.ok_button = tk.Button(self.ok_button_frame, text="OK", state="disabled", command=self.save_key_and_proceed, fg="cyan2", bg="black", relief="flat")
        self.ok_button.pack(padx=1, pady=1)
        self.key_entry.bind("<KeyRelease>", self.check_key) # Check key on release

    def on_entry_click(self, event):
        if self.key_entry.get() == "Paste your key here.":
            self.key_entry.delete(0, "end")
            self.key_entry.config(fg="white") # Change text color when focused

    def on_focus_out(self, event):
        if not self.key_entry.get(): # Check if empty
            self.key_entry.insert(0, "Paste your key here.")
            self.key_entry.config(fg="dark grey") # Restore placeholder color

    def check_key(self, event):
        """Enable OK button only if the key entry is not empty and not the placeholder."""
        key = self.key_entry.get()
        if key and key != "Paste your key here.":
            self.ok_button.config(state="normal")
        else:
            self.ok_button.config(state="disabled")


    def save_key_and_proceed(self):
        key = self.key_entry.get()
        if key and key != "Paste your key here.":
            env_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), ".env")
            try:
                lines = []
                token_found = False
                if os.path.exists(env_path):
                    with open(env_path, "r") as f:
                        lines = f.readlines()

                with open(env_path, "w") as f:
                    for line in lines:
                        if line.strip().startswith("TOKEN="):
                            f.write(f"TOKEN={key}")
                            token_found = True
                        else:
                            f.write(line)
                    if not token_found: # If TOKEN line didn't exist, append it
                         f.write(f"TOKEN={key}")

                messagebox.showinfo("Success", "Key saved successfully")
                # Decide where to go next, perhaps back to choice or to processing?
                # Going back to Choice frame for now.
                self.controller.show_frame(self.controller.choice_frame)
            except Exception as e:
                messagebox.showerror("Error", f"Failed to save key to {env_path}: {str(e)}")
        else:
            messagebox.showerror("Error", "Please enter a valid key")


### Processing Frame (Window 2) ###
# Helper class to redirect stdout/stderr to a Tkinter Text widget
class TextRedirector(io.StringIO):
    def __init__(self, widget):
        super().__init__()
        self.widget = widget

    def write(self, s):
        # Ensure GUI updates happen in the main thread
        self.widget.after(0, self._write_to_widget, s)

    def _write_to_widget(self, s):
        try:
             if self.widget.winfo_exists(): # Check if widget still exists
                 self.widget.config(state=tk.NORMAL)
                 self.widget.insert(tk.END, s)
                 self.widget.see(tk.END) # Scroll to the end
                 self.widget.config(state=tk.DISABLED)
        except tk.TclError as e:
             # Handle cases where widget might be destroyed during async update
             print(f"TclError updating text widget: {e}")


class ProcessingFrame(tk.Frame):
    def __init__(self, parent, controller):
        tk.Frame.__init__(self, parent, bg="gray8")
        self.controller = controller
        self.sd_checkbuttons = [] # To store references to checkbuttons
        self.other_buttons = [] # To store refs to Back, Proceed, Browse buttons

        # --- Scrollbar Setup ---
        # Main frame for content, allows us to control overall padding/centering if needed later
        main_content_frame = tk.Frame(self, bg="gray8")
        main_content_frame.pack(side="top", fill="both", expand=True)

        self.canvas = tk.Canvas(main_content_frame, bg="gray8", highlightthickness=0)
        # Make scrollbar part of main_content_frame to appear next to canvas
        self.scrollbar = ttk.Scrollbar(main_content_frame, orient="vertical", command=self.canvas.yview)
        # The scrollable frame goes INSIDE the canvas
        self.scrollable_frame = tk.Frame(self.canvas, bg="gray8")

        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: self.canvas.configure(
                scrollregion=self.canvas.bbox("all"),
                # Optional: Set width of scrollable frame to canvas width
                # width=e.width
            )
        )

        # Add the scrollable frame to the canvas
        self.canvas_window = self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        self.canvas.configure(yscrollcommand=self.scrollbar.set)

        # Pack canvas and scrollbar side-by-side within main_content_frame
        self.scrollbar.pack(side="right", fill="y")
        self.canvas.pack(side="left", fill="both", expand=True)

        # Re-bind configure event to ensure canvas window width matches canvas width
        self.canvas.bind('<Configure>', self._on_canvas_configure)
        # --- End Scrollbar Setup ---


        # --- Widgets go into self.scrollable_frame now ---
        # Back button - Keep anchored to top-left
        back_frame = tk.Frame(self.scrollable_frame, bg="dark slate gray")
        # Use pack directly on scrollable_frame, not a separate frame for positioning
        back_frame.pack(anchor="nw", padx=10, pady=10)
        back_button = tk.Button(back_frame, text="Back", command=lambda: controller.show_frame(controller.choice_frame), fg="dark slate gray", bg="black", relief="flat")
        back_button.pack(padx=1, pady=1)
        self.other_buttons.append(back_button) # Store ref


        # --- Centered Content Frame ---
        # Create a frame within scrollable_frame to hold all centered content
        center_frame = tk.Frame(self.scrollable_frame, bg="gray8")
        # Pack the center_frame itself, it will expand horizontally
        center_frame.pack(pady=10, padx=20, fill='x', expand=True)


        # --- Widgets to be centered go inside center_frame ---

        # TOE type selection
        toe_label = tk.Label(center_frame, text="Please select the TOE type:", fg="white", bg="gray8", font=("Arial", 13, "bold"))
        # Pack widgets inside center_frame, they will be centered by default pack behavior within this frame
        toe_label.pack(pady=(10, 5))
        options = ["STANDALONE", "DISTRIBUTED"]
        self.combo = ttk.Combobox(center_frame, values=options, state="readonly", width=25, font=("Arial", 10)) # Added font, width
        self.combo.pack(pady=5)
        self.other_buttons.append(self.combo) # Store ref

        # Generate files button ("Proceed")
        generate_frame = tk.Frame(center_frame, bg="cyan2")
        generate_frame.pack(pady=10) # This frame will be centered
        generate_button = tk.Button(generate_frame, font=("Arial", 10), text="Proceed", command=self.generate_files, fg="cyan2", bg="black", relief="flat", width=15) # Added width
        generate_button.pack(padx=1, pady=1)
        self.other_buttons.append(generate_button) # Store ref

        # ST selection
        st_label = tk.Label(center_frame, text="Select ST document:", fg="white", bg="gray8", font=("Arial", 13, "bold"))
        st_label.pack(pady=(15, 5))
        self.st_path_var = tk.StringVar()
        st_button_frame = tk.Frame(center_frame, bg="cyan2")
        st_button_frame.pack(pady=5) # Centered
        self.st_button = tk.Button(st_button_frame, font=("Arial", 10), text="Browse", command=self.browse_st, fg="cyan2", bg="black", relief="flat", state="disabled", width=15) # Added width
        self.st_button.pack(padx=1, pady=1)
        self.other_buttons.append(self.st_button) # Store ref
        self.st_selected_label = tk.Label(center_frame, text="No ST selected.", fg="dark grey", bg="gray8", font=("Arial", 10))
        self.st_selected_label.pack(pady=(5, 15))

        # SD selection Label
        sd_label = tk.Label(center_frame, text="Please select the PP/MODS/PKGs for this project's evaluation:", fg="white", bg="gray8", font=("Arial", 13, "bold"))
        sd_label.pack(pady=10)

        # --- SD LabelFrames (These should still fill width relative to center_frame) ---
        self.sd_vars = {}
        style = ttk.Style()
        style.configure("Black.TLabelframe", background="gray8", bordercolor="gray50")
        style.configure("Black.TLabelframe.Label", foreground="cyan2", background="gray8", font=("Arial", 12, "bold"))

        for category, options_dict in SD_OPTIONS.items():
            # Pack the LabelFrames within center_frame, but make them fill 'x'
            frame = ttk.LabelFrame(center_frame, text=category, style="Black.TLabelframe")
            # fill='x' makes the LabelFrame expand horizontally within center_frame
            frame.pack(fill="x", padx=10, pady=5, expand=True) # Added expand=True
            for option in options_dict:
                var = tk.BooleanVar()
                self.sd_vars[option] = var
                cb = tk.Checkbutton(frame, text=option, variable=var, fg="cyan2", bg="gray8", selectcolor="black", state="disabled", anchor="w", justify="left", font=("Arial", 10)) # Added font
                # Pack checkbuttons inside the LabelFrame
                cb.pack(anchor="w", padx=20, pady=3) # Increased padx for indentation
                self.sd_checkbuttons.append(cb) # Store checkbutton references

        # Validate TSS button
        process_frame = tk.Frame(center_frame, bg="cyan2")
        process_frame.pack(pady=20) # Centered
        self.process_button = tk.Button(process_frame, font=("Arial", 10), text="Validate TSS", command=self.process_requirements, fg="cyan2", bg="black", relief="flat", state="disabled", width=15) # Added width
        self.process_button.pack(padx=3, pady=3)

        # Progress bar and message Frame (also centered)
        self.progress_outer_frame = tk.Frame(center_frame, bg="gray8")
        # Don't pack initially, pack when needed
        self.progress_label = tk.Label(self.progress_outer_frame, text="Processing... This may take a while.\n\nConsole output below:", fg="white", bg="gray8", font=("Arial", 12))
        self.progress_label.pack(pady=(5,0))
        self.progress_bar = ttk.Progressbar(self.progress_outer_frame, mode="indeterminate")
        # Make progress bar fill horizontally within its frame
        self.progress_bar.pack(pady=5, fill='x', padx=20, expand=True)

        # Console Output Text Area
        self.console_output = tk.Text(self.progress_outer_frame, height=10, state="disabled", bg="black", fg="lightgrey", font=("Consolas", 9), wrap=tk.WORD, relief="sunken", bd=1) # Added relief/bd
        # Make console fill horizontally
        self.console_output.pack(pady=5, padx=20, fill='x', expand=True)

        # --- End of centered content ---

        # Bind mouse wheel for scrolling on canvas and scrollable frame
        # Binding to self (ProcessingFrame) should be sufficient if focus is managed
        self.bind_all("<MouseWheel>", self._on_mousewheel)
        # self.canvas.bind_all("<MouseWheel>", self._on_mousewheel) # Might not be needed
        # self.scrollable_frame.bind_all("<MouseWheel>", self._on_mousewheel) # Might not be needed

    def _on_canvas_configure(self, event):
        """Reset the scrollable frame's width to the canvas width."""
        canvas_width = event.width
        self.canvas.itemconfig(self.canvas_window, width=canvas_width)
        self.canvas.configure(scrollregion=self.canvas.bbox("all"))

    def _on_mousewheel(self, event):
        # Determine scroll direction (platform dependent)
        # Check if the mouse pointer is within the canvas area
        canvas_x = self.canvas.winfo_rootx()
        canvas_y = self.canvas.winfo_rooty()
        canvas_w = self.canvas.winfo_width()
        canvas_h = self.canvas.winfo_height()

        # Check if the event coordinates are within the canvas bounds
        if (canvas_x <= event.x_root < canvas_x + canvas_w and
            canvas_y <= event.y_root < canvas_y + canvas_h):
            delta = 0
            if sys.platform.startswith('win') or sys.platform.startswith('darwin'): # Windows or Mac OS
                 delta = -1 * int(event.delta / 120)
            elif event.num == 4: # Linux scroll up
                 delta = -1
            elif event.num == 5: # Linux scroll down
                 delta = 1

            if delta:
                 self.canvas.yview_scroll(delta, "units")


    def browse_st(self):
        file_path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
        if file_path:
            self.st_path_var.set(file_path)
            self.st_selected_label.config(text=f"Selected ST: {os.path.basename(file_path)}", fg="white") # Change color on selection
        else:
            self.st_path_var.set("") # Clear path if cancelled
            self.st_selected_label.config(text="No ST selected.", fg="dark grey") # Reset text and color


    def generate_files(self):
        selection = self.combo.get()
        if not selection:
            messagebox.showerror("Error", "Please select a TOE type from the dropdown.")
            return

        # Create sys_inst directory if it doesn't exist
        sys_inst_dir = os.path.dirname(OUTPUT_TSS_PATH)
        os.makedirs(sys_inst_dir, exist_ok=True)


        try:
            process_file(BASE_TSS_PATH, OUTPUT_TSS_PATH, selection)
             # No success message box, rely on print statement and enabling buttons
            self.st_button.config(state="normal")
            for cb in self.sd_checkbuttons:
                 cb.config(state="normal")
            self.process_button.config(state="normal")
            # Optionally disable the "Proceed" button itself after success?
            # self.other_buttons[2].config(state="disabled") # Index 2 is generate_button
        except FileNotFoundError:
             # Error already shown by process_file
             return
        except Exception as e:
             # Error already shown by process_file
             return


    def set_buttons_state(self, state):
        """Enable or disable relevant buttons."""
        # Disable/Enable Back, Proceed, Browse, Combobox
        for btn in self.other_buttons:
             try: # Use try-except in case a widget reference is broken
                 btn.config(state=state)
             except tk.TclError:
                 print(f"Warning: Could not configure widget {btn}")


        # Disable/Enable Checkbuttons
        for cb in self.sd_checkbuttons:
            try:
                cb.config(state=state)
            except tk.TclError:
                 print(f"Warning: Could not configure widget {cb}")

        # Special handling for Process button (disable when processing starts, enable when done/error)
        # The Validate TSS button itself should be disabled during processing
        if state == "disabled":
             self.process_button.config(state="disabled")
        # Re-enabling Process button depends on whether ST/SDs are selected after processing finishes
        # We'll handle its re-enabling specifically in run_processing's finally block or processing_done


    def process_requirements(self):
        st_path = self.st_path_var.get()
        if not st_path:
            messagebox.showerror("Error", "Please select an ST document")
            return
        if not os.path.exists(st_path):
            messagebox.showerror("Error", f"ST document not found: {st_path}")
            return


        selected_sd = {option: SD_OPTIONS[category][option]
                       for category in SD_OPTIONS
                       for option in SD_OPTIONS[category]
                       if option in self.sd_vars and self.sd_vars[option].get()}

        if not selected_sd:
            messagebox.showerror("Error", "Please select at least one PP/MOD/PKG document")
            return

        selected_sd_paths = list(selected_sd.values())

        # --- Start Processing ---
        self.progress_outer_frame.pack(pady=10, fill='x', padx=10) # Show progress section
        self.progress_bar.start()
        self.set_buttons_state("disabled") # Disable buttons

         # Clear previous console output
        self.console_output.config(state=tk.NORMAL)
        self.console_output.delete('1.0', tk.END)
        self.console_output.config(state=tk.DISABLED)


        # --- Redirect stdout/stderr ---
        self.stdout_redirector = TextRedirector(self.console_output)
        self.stderr_redirector = TextRedirector(self.console_output)
        self.original_stdout = sys.stdout
        self.original_stderr = sys.stderr
        sys.stdout = self.stdout_redirector
        sys.stderr = self.stderr_redirector
        # --- End Redirection ---


        print("--- Starting Validation Process ---")
        print(f"ST Document: {os.path.basename(st_path)}")
        print(f"Selected SDs: {', '.join(selected_sd.keys())}")

        # Run processing in a separate thread
        thread = threading.Thread(target=self.run_processing, args=(st_path, selected_sd_paths), daemon=True)
        thread.start()

    def run_processing(self, st_path, sd_paths):
        processing_success = False
        final_message = "An unexpected error occurred during processing."
        try:
            processor = RequirementsProcessor()
            print("\n--- Running TSS/SFR/SD Extraction ---")
            success, message = processor.process_files(st_path, sd_paths)

            if not success:
                 # Error message already printed by process_files or extract_st_data
                 # Show error in messagebox from main thread
                 self.controller.after(0, lambda msg=message: messagebox.showerror("Processing Error", msg))
                 final_message = message # Store for finally block
                 return # Stop processing further

            print("\n--- TSS/SFR/SD Extraction Successful ---")
            print(message) # Print success message to console


            # --- Subprocess Execution ---
            scripts_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Scripts")
            python_executable = os.path.join(scripts_dir, "python.exe") # Assuming python.exe is directly in Scripts
            # Check if the custom python executable exists
            if not os.path.exists(python_executable):
                 # Fallback to system's Python if custom one isn't found
                 python_executable = sys.executable
                 print(f"\nWarning: Custom Python not found at {os.path.join(scripts_dir, 'python.exe')}. Falling back to system Python: {python_executable}")


            try:
                # --- Trigger the API processing subprocess ---
                api_pros_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "api_processing_deb.py")
                if os.path.exists(api_pros_path):
                    print(f"\n--- Running API Processing Subprocess ({os.path.basename(api_pros_path)}) ---")
                    # Use capture_output=True to get stdout/stderr from subprocess
                    result_api = subprocess.run([python_executable, api_pros_path], check=True, capture_output=True, text=True, encoding='utf-8')
                    print("--- API Subprocess Output ---")
                    print(result_api.stdout)
                    if result_api.stderr:
                        print("--- API Subprocess Errors ---")
                        print(result_api.stderr)
                    print("--- API Processing Subprocess Completed ---")
                else:
                    print(f"\nWarning: API processing script not found at {api_pros_path}. Skipping.")

                # --- Call AARF.py with selected SD names ---
                selected_sd_names = [opt for category in SD_OPTIONS for opt in SD_OPTIONS[category] if opt in self.sd_vars and self.sd_vars[opt].get()]
                aarf_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "AARF.py")
                if os.path.exists(aarf_path):
                    print(f"\n--- Running AARF Subprocess ({os.path.basename(aarf_path)}) ---")
                    print(f"Arguments: {selected_sd_names}")
                    result_aarf = subprocess.run([python_executable, aarf_path] + selected_sd_names, check=True, capture_output=True, text=True, encoding='utf-8')
                    print("--- AARF Subprocess Output ---")
                    print(result_aarf.stdout)
                    if result_aarf.stderr:
                        print("--- AARF Subprocess Errors ---")
                        print(result_aarf.stderr)
                    print("--- AARF Subprocess Completed ---")
                else:
                    print(f"\nWarning: AARF script not found at {aarf_path}. Skipping.")

                processing_success = True # Mark overall success
                final_message = "Validation for TSS has been completed successfully!\nResults have been stored to the 'BLITZ-output' folder.\n\nNote: If TOE is distributed then please append the 'General Requirements for Distributed TOE' section to the AAR."


            except subprocess.CalledProcessError as e:
                # Error in subprocess execution
                script_name = os.path.basename(e.cmd[1]) if len(e.cmd) > 1 else "Subprocess"
                error_detail = f"Subprocess '{script_name}' failed with return code {e.returncode}."
                print(f"\n--- Subprocess Error ---")
                print(error_detail)
                if e.stdout:
                    print("--- Subprocess Output (stdout) ---")
                    print(e.stdout)
                if e.stderr:
                    print("--- Subprocess Error Output (stderr) ---")
                    print(e.stderr)
                final_message = f"{error_detail}\nSee console for details."
                self.controller.after(0, lambda msg=final_message: messagebox.showerror("Subprocess Error", msg))
                return # Stop processing
            except FileNotFoundError as e:
                 # Could happen if python executable is wrong
                 error_detail = f"Failed to find executable for subprocess: {e}"
                 print(f"\n--- Subprocess Error ---")
                 print(error_detail)
                 final_message = error_detail
                 self.controller.after(0, lambda msg=final_message: messagebox.showerror("Subprocess Error", msg))
                 return # Stop processing

            # --- Processing Done (if successful) ---
            # Schedule the final success message box from the main thread
            self.controller.after(0, self.processing_done, final_message)

        except Exception as e:
            # Catch any other unexpected errors during the thread's execution
            print(f"\n--- Unexpected Error in Processing Thread ---")
            import traceback
            print(traceback.format_exc()) # Print detailed traceback to console
            final_message = f"An unexpected error occurred: {str(e)}\nSee console for details."
            # Show error in messagebox from main thread
            self.controller.after(0, lambda msg=final_message: messagebox.showerror("Unexpected Error", msg))

        finally:
            # --- Restore stdout/stderr ---
            sys.stdout = self.original_stdout
            sys.stderr = self.original_stderr
            # --- Stop progress bar and re-enable buttons ---
            # Ensure these GUI updates run in the main thread using 'after'
            self.controller.after(0, self.finalize_processing_ui, processing_success)


    def finalize_processing_ui(self, was_successful):
        """Handles stopping progress bar and re-enabling buttons in the main thread."""
        self.progress_bar.stop()
        # Don't hide the progress frame immediately, user might want to read console
        # self.progress_outer_frame.pack_forget()
        print("\n--- Validation Process Finished ---")
        self.set_buttons_state("normal") # Re-enable most buttons

        # Decide final state of process_button and ST/SD selectors
        st_selected = bool(self.st_path_var.get())
        sd_selected = any(var.get() for var in self.sd_vars.values())

        if st_selected and sd_selected:
            self.process_button.config(state="normal")
        else:
            self.process_button.config(state="disabled")

        # Keep ST button enabled if Proceed was successful (user might want to change ST)
        # Keep SD checkbuttons enabled too
        # Keep Proceed button enabled (user might want to change TOE type and proceed again)


    def processing_done(self, message):
        """Called from main thread after successful processing."""
        messagebox.showinfo("Success", message)
        # Decide if you want to close the app automatically on success
        # self.controller.destroy() # Uncomment to close app on success

### Main Execution
if __name__ == "__main__":
    # Make sure required directories exist
    os.makedirs(os.path.join(os.path.dirname(os.path.abspath(__file__)), "sys_inst"), exist_ok=True)
    os.makedirs(os.path.join(os.path.dirname(os.path.abspath(__file__)), "DB"), exist_ok=True)
    os.makedirs(os.path.join(os.path.dirname(os.path.abspath(__file__)), "UI"), exist_ok=True)
    os.makedirs(os.path.join(os.path.dirname(os.path.abspath(__file__)), "ephemeral"), exist_ok=True)
    os.makedirs(os.path.join(os.path.dirname(os.path.abspath(__file__)), "BLITZ-output"), exist_ok=True)
    # Optional: Create dummy .env if it doesn't exist
    env_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), ".env")
    if not os.path.exists(env_file):
        with open(env_file, "w") as f:
            f.write("TOKEN=")


    app = BlitzApp()
    app.mainloop()