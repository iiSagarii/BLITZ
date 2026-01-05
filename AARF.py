import sys
import os
import json
import docx
import docx.oxml
from docx import Document
from docx.oxml import parse_xml
from lxml import etree
from openpyxl import Workbook
from openpyxl.styles import Font
from docx.oxml import CT_P, CT_Tbl
from docx.shared import Pt, RGBColor
import re

# Hardcoded paths (Update these paths as needed)
TEMPLATE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "templates") # points towards the templates 
JSON_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "ephemeral/ai_responses.json") #points towards ai_responses
OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "BLITZ-output")

# Ensure output directory exists
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Get selected SDs from command-line arguments or use default
selected_sds = sys.argv[1:] if len(sys.argv) > 1 else ["NDcPP_v3.0"] # Example: ["NDcPP_v3.0", "PKG_SSH_v1.0"]
#selected_sds = ["NDcPP_v3.0"] # Default for testing

print(f"Processing SDs: {selected_sds}")
template_paths = [os.path.join(TEMPLATE_DIR, f"{sd}-template.docx") for sd in selected_sds]

# Load JSON input
try:
    
    # Load from JSON file 
    with open(JSON_PATH, 'r', encoding='utf-8') as f:
        data = json.load(f)
    # -------------------------
    doc_data = data.get("DOC", [])
    excel_data = data.get("Excel", [])
    print(f"Loaded JSON data.") #(Source: {JSON_PATH if 'f' in locals() else 'string'})")
except FileNotFoundError:
    print(f"Error: JSON input file not found at {JSON_PATH}")
    sys.exit(1)
except json.JSONDecodeError as e:
    print(f"Error: Could not decode JSON: {e}")
    sys.exit(1)

# Load all template documents
template_docs = []
for path in template_paths:
    if os.path.exists(path):
        try:
            template_docs.append(Document(path))
            print(f"Loaded template: {path}")
        except Exception as e:
            print(f"Error loading template {path}: {e}")
    else:
        print(f"Warning: Template file not found: {path}")

if not template_docs:
    print("Error: No template documents were successfully loaded. Exiting.")
    sys.exit(1)

# Helper function to iterate through block items (paragraphs and tables) in order
def iter_block_items(parent):
    if isinstance(parent, docx.document.Document):
        parent_elm = parent.element.body
    elif isinstance(parent, docx.table._Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("Unsupported parent type")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield docx.text.paragraph.Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield docx.table.Table(child, parent)

# HELPER FUNCTION to find placeholders
def find_placeholders(text):
    """Finds all placeholders like <Ans#...> in a string."""
    # Regex to find <Ans# followed by one or more digits, then >
    return re.findall(r"<Ans#\d+>", text)

# Build hierarchical heading structure for a document
def build_heading_structure(doc):
    structure = []
    current_h3 = None
    current_h4 = None
    current_h5 = None
    for block in iter_block_items(doc):
        if isinstance(block, docx.text.paragraph.Paragraph):
            style_name = block.style.name
            text = block.text.strip()

            if style_name == 'Heading 3':
                current_h3 = {'paragraph': block, 'subheadings': [], 'needed': False}
                structure.append(current_h3)
                current_h4 = None
                current_h5 = None
            elif style_name == 'Heading 4':
                if current_h3 is not None:
                    current_h4 = {'paragraph': block, 'subheadings': [], 'needed': False}
                    current_h3['subheadings'].append(current_h4)
                    current_h5 = None
            elif style_name == 'Heading 5':
                if current_h4 is not None:
                    is_tss = text.endswith(" TSS")
                    is_agd = text.endswith(" AGD")
                    sfr_base = text[:-4] if is_tss or is_agd else text # Base name
                    current_h5 = {
                        'paragraph': block,
                        'content': [],
                        'referenced': False,
                        'is_tss': is_tss,
                        'is_agd': is_agd,
                        'agd_content': [],
                        'linked_as_agd': False,
                        'sfr_base': sfr_base
                    }
                    current_h4['subheadings'].append(current_h5)
            else: # Content under the current Heading 5
                if current_h5 is not None:
                    current_h5['content'].append(block)
        elif isinstance(block, docx.table.Table):
            if current_h5 is not None:
                current_h5['content'].append(block)

    return structure

# Link AGD sections to their corresponding TSS sections
def link_agd_to_tss(structure):
    for h3 in structure:
        for h4 in h3['subheadings']:
            h5_list = h4['subheadings']
            for i, h5_node in enumerate(h5_list):
                if h5_node['is_tss'] and not h5_node['agd_content']:
                    tss_sfr_base = h5_node['sfr_base']
                    if i + 1 < len(h5_list):
                        next_h5_node = h5_list[i+1]
                        if next_h5_node['is_agd'] and next_h5_node['sfr_base'] == tss_sfr_base and not next_h5_node['linked_as_agd']:
                            print(f"Linking AGD '{next_h5_node['paragraph'].text.strip()}' to TSS '{h5_node['paragraph'].text.strip()}'")
                            h5_node['agd_content'].append(next_h5_node['paragraph'])
                            h5_node['agd_content'].extend(next_h5_node['content'])
                            next_h5_node['linked_as_agd'] = True

# Replace placeholders in the second row of the table using the new multi-replace function
def modify_table(table, answers):
    """Modifies the second row of the table, replacing placeholders in each cell.

    Args:
        table: The docx table object.
        answers: A dictionary of placeholders (e.g., '<Ans#1>') and their values for the current SFR.
    """
    if len(table.rows) >= 2:
        second_row = table.rows[1]
        for cell_idx, cell in enumerate(second_row.cells):
            print(f"Processing Cell {cell_idx} in row 1: '{cell.text[:50]}...'") # Debug print
            # Apply all replacements to each paragraph within the cell
            for paragraph in cell.paragraphs:
                 # Pass the paragraph and the *entire* answers dict to the new function
                 replace_all_placeholders_in_paragraph(paragraph, answers)
    else:
        print("Warning: Table has less than 2 rows, cannot modify.")


# Replace text in a cell's paragraphs, handling splits across runs. Replace multiple placeholders within a single paragraph's text
# Replace multiple placeholders and return True if any changes were made
def replace_all_placeholders_in_paragraph(paragraph, replacements):
    """Replaces all placeholders found in the paragraph's text using the replacements dictionary.

    Args:
        paragraph: The docx paragraph object.
        replacements: A dictionary where keys are placeholders (e.g., '<Ans#1>')
                      and values are the replacement strings.

    Returns:
        bool: True if any placeholder was found and replaced in the paragraph, False otherwise.
    """
    original_text = paragraph.text
    modified_text = original_text
    placeholders_replaced = False # Renamed for clarity

    # Iterate through the provided replacements and apply them to the text
    for placeholder, value in replacements.items():
        placeholder_pattern = f"<{placeholder}>"
        if placeholder in modified_text:
            # print(f"    Found '{placeholder}' in paragraph. Replacing.") # Debug print
            modified_text = modified_text.replace(placeholder_pattern, str(value)) # Ensure value is string
            placeholders_replaced = True # Mark that at least one replacement happened

    # If any replacements were made, rebuild the paragraph runs
    if placeholders_replaced:
        # Clear existing runs
        p_element = paragraph._p
        runs_to_remove = list(p_element.r_lst)
        for run_element in runs_to_remove:
            p_element.remove(run_element)

        # Add a new run with the fully modified text
        new_run = paragraph.add_run(modified_text)
        new_run.bold = True
        # print(f"    Rebuilt paragraph with modified text.") # Debug print
        return True # Indicate replacement occurred
    else:
        # print(f"    No placeholders from replacements dict found in paragraph: '{original_text[:60]}...'") # Debug print
        return False # Indicate no replacement occurred


# Copy a block (paragraph or table) to the target document
def copy_block_to_doc(target_doc, block):
    if isinstance(block, docx.text.paragraph.Paragraph):
        elem = block._element
    elif isinstance(block, docx.table.Table):
        elem = block._tbl
    else:
        return

    try:
        if elem is not None:
            xml_string = etree.tostring(elem, encoding='unicode')
            if xml_string:
                copied_elem = parse_xml(xml_string)
                target_doc.element.body.append(copied_elem)
            # else: print(f"Warning: XML string for {type(block).__name__} is empty.")
        # else: print(f"Warning: Element for {type(block).__name__} is None.")
    except Exception as e:
        print(f"Error copying element {type(block).__name__}: {e}")


# --- Main Script Logic ---

# Build structures and link AGD
structures = {}
for sd, doc in zip(selected_sds, template_docs):
    print(f"\nBuilding structure for {sd}...")
    structures[sd] = build_heading_structure(doc)
    print(f"Linking AGD sections for {sd}...")
    link_agd_to_tss(structures[sd])

# Build mapping from base SFR name to (sd, h5_node) for TSS sections
sfr_to_tss_node = {}
for sd, structure in structures.items():
    for h3 in structure:
        for h4 in h3['subheadings']:
            for h5 in h4['subheadings']:
                if h5['is_tss']:
                    sfr_base = h5['sfr_base']
                    if sfr_base not in sfr_to_tss_node:
                       sfr_to_tss_node[sfr_base] = (sd, h5)

# Mark referenced TSS nodes and modify their content based on JSON
print("\nProcessing JSON data for document modifications...")
for obj in doc_data:
    sfr = obj.get("SFR")
    if sfr and sfr in sfr_to_tss_node:
        sd, h5_node = sfr_to_tss_node[sfr]
        print(f"\nProcessing SFR: {sfr} (found in {sd})") # Debug print

        # Mark node as referenced if not already
        if not h5_node['referenced']:
            h5_node['referenced'] = True
            print(f"  Marked TSS as referenced: '{h5_node['paragraph'].text.strip()}'")

        # Extract answers and store the keys provided for this SFR
        answers = {k: v for k, v in obj.items() if k.startswith("Ans#")}
        provided_answer_keys = set(answers.keys()) # Store keys like {'Ans#1', 'Ans#2', 'Ans#4'}
        h5_node['provided_answer_keys'] = provided_answer_keys # Attach to node
        print(f"  Provided answer keys for {sfr}: {provided_answer_keys}")

        if not answers:
            print(f"  No 'Ans#' found for SFR '{sfr}' in JSON object. Skipping content modification.") # Debug print
            continue # No answers to process for this SFR

        print(f"  Applying replacements to content blocks of '{h5_node['paragraph'].text.strip()}'...") # Debug print
        # Iterate through ALL content blocks (paragraphs and tables) associated with this H5
        for block_idx, block in enumerate(h5_node['content']):
            if isinstance(block, docx.text.paragraph.Paragraph):
                # Apply replacements directly to the paragraph
                # print(f"    Processing content paragraph {block_idx}: '{block.text[:50]}...'") # Debug
                replace_all_placeholders_in_paragraph(block, answers)
            elif isinstance(block, docx.table.Table):
                # Apply replacements to all paragraphs within all cells of the table
                # print(f"    Processing content table {block_idx}...") # Debug
                for row_idx, row in enumerate(block.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for para_idx, paragraph in enumerate(cell.paragraphs):
                            # print(f"      Processing table para ({row_idx},{cell_idx},{para_idx}): '{paragraph.text[:50]}...'") # Debug
                            replace_all_placeholders_in_paragraph(paragraph, answers)
    elif sfr:
        print(f"Warning: SFR '{sfr}' from JSON not found as a TSS heading in templates.")

# Propagate 'needed' flags upwards
print("\nPropagating 'needed' flags...")
for sd, structure in structures.items():
    for h3 in structure:
        h3_needed = False
        for h4 in h3['subheadings']:
            h4_needed = False
            for h5 in h4['subheadings']:
                is_needed_agd = h5['linked_as_agd'] and any(tss_h5['referenced'] for tss_h5 in h4['subheadings'] if tss_h5['is_tss'] and tss_h5['sfr_base'] == h5['sfr_base'])
                if h5['referenced'] or is_needed_agd:
                    h4_needed = True
                    h3_needed = True
            h4['needed'] = h4_needed
        h3['needed'] = h3_needed

# Create the final AAR-TSS document
print("\nCreating final AAR-TSS document...")
final_doc = Document()

# Apply basic styles (or copy from template if needed)
try:
    heading3_style = final_doc.styles['Heading 3']
    # Apply formatting as needed
    heading4_style = final_doc.styles['Heading 4']
    # Apply formatting as needed
    heading5_style = final_doc.styles['Heading 5']
    # Apply formatting as needed
    normal_style = final_doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
except KeyError as e:
    print(f"Warning: Style {e} not found in default document. Formatting may differ.")

# Add Content to Final Document (including General Reqs and Needed sections)
general_req_heading_part = "GENERAL REQUIREMENTS FOR DISTRIBUTED TOES"
general_req_added = False


# --- Add Content to Final Document ---

# Add the general requirements section first, if it exists and is needed
general_req_heading_part = "GENERAL REQUIREMENTS FOR DISTRIBUTED TOES" # Match partial heading
general_req_added = False
for sd in selected_sds:
    structure = structures[sd]
    for h3 in structure:
        if h3['paragraph'].text.strip().startswith(general_req_heading_part):
             print(f"Adding General Requirements section found in {sd}")
             copy_block_to_doc(final_doc, h3['paragraph'])
             for h4 in h3['subheadings']:
                 copy_block_to_doc(final_doc, h4['paragraph'])
                 for h5 in h4['subheadings']:
                     copy_block_to_doc(final_doc, h5['paragraph'])
                     for block in h5['content']:
                         copy_block_to_doc(final_doc, block)
                     if h5['agd_content']:
                          for agd_block in h5['agd_content']:
                              copy_block_to_doc(final_doc, agd_block)
             general_req_added = True
             break
    if general_req_added:
        break

if not general_req_added:
    print("General Requirements section not found or not added.")

# Add all needed sections based on referenced TSS and their parents/linked AGDs
processed_h5_texts = set()

# Add all needed sections based on referenced TSS and their parents/linked AGDs
processed_h5_texts = set() # Keep track of H5 headings already added to avoid duplication

for sd in selected_sds:
    structure = structures[sd]
    for h3 in structure:
        # Skip the general reqs H3 if it was handled separately and already added
        if general_req_added and h3['paragraph'].text.strip().startswith(general_req_heading_part):
            continue

        if h3['needed']:
            print(f"Adding needed H3: {h3['paragraph'].text.strip()}")
            copy_block_to_doc(final_doc, h3['paragraph']) # Copy H3 heading

            for h4 in h3['subheadings']:
                if h4['needed']:
                    print(f"  Adding needed H4: {h4['paragraph'].text.strip()}")
                    copy_block_to_doc(final_doc, h4['paragraph']) # Copy H4 heading

                    for h5 in h4['subheadings']:
                        h5_text = h5['paragraph'].text.strip()
                        is_agd_linked_to_needed_tss = h5['linked_as_agd'] # Check if it's an AGD linked earlier

                        # Primary condition: Process H5 if it's a TSS node that was referenced by JSON
                        if h5['referenced'] and not h5['linked_as_agd']:
                           if h5_text not in processed_h5_texts:
                                print(f"    Adding referenced TSS H5: {h5_text}")
                                copy_block_to_doc(final_doc, h5['paragraph']) # Copy H5 heading
                                processed_h5_texts.add(h5_text) # Mark heading as added

                                # --- Conditional Content Copying ---
                                # Get the set of answer keys provided for this specific H5 node
                                provided_keys = h5.get('provided_answer_keys', set()) # Default to empty set
                                print(f"      Processing {len(h5['content'])} content block(s) for {h5_text}. Provided keys: {provided_keys}") # Debug

                                for block_idx, block in enumerate(h5['content']):
                                    copy_this_block = True # Assume we copy unless a rule says otherwise
                                    block_text = ""

                                    # Get text content to check for placeholders
                                    if isinstance(block, docx.text.paragraph.Paragraph):
                                        block_text = block.text
                                    elif isinstance(block, docx.table.Table):
                                        # Combine text from all cells for placeholder checking
                                        try:
                                            block_text = "\n".join(cell.text for row in block.rows for cell in row.cells)
                                        except Exception:
                                             block_text = "" # Handle potential errors getting table text

                                    # Find placeholders in the block's text (using the helper function)
                                    placeholders_in_block = find_placeholders(block_text)

                                    if placeholders_in_block:
                                        # If there are placeholders, check if the *first* one found was provided in the JSON
                                        primary_placeholder = placeholders_in_block[0] # e.g., "<Ans#3>"
                                        # Extract the key part, e.g., "Ans#3"
                                        primary_key = primary_placeholder.strip("<>") # Removes < and >

                                        if primary_key not in provided_keys:
                                            copy_this_block = False # Skip this block
                                            print(f"        Skipping block {block_idx} ({type(block).__name__}) because primary placeholder '{primary_placeholder}' key '{primary_key}' not in provided keys {provided_keys}") # Debug

                                    # Copy the block if it passed the check
                                    if copy_this_block:
                                        # print(f"        Copying block {block_idx} ({type(block).__name__})") # Debug
                                        copy_block_to_doc(final_doc, block)
                                    # else: Already printed reason for skipping
                                # --- End Conditional Content Copying ---

                                # --- Add linked AGD content unconditionally if TSS was added ---
                                if h5['agd_content']:
                                    agd_h5_para = h5['agd_content'][0] # First block is the AGD H5 paragraph
                                    agd_h5_text = agd_h5_para.text.strip()
                                    print(f"    Adding linked AGD section: {agd_h5_text}")
                                    if agd_h5_text not in processed_h5_texts:
                                        # Copy all blocks stored in agd_content (heading + content)
                                        for agd_block in h5['agd_content']:
                                            copy_block_to_doc(final_doc, agd_block)
                                        processed_h5_texts.add(agd_h5_text) # Mark AGD heading as processed
                                    else:
                                         print(f"    Skipping already processed AGD: {agd_h5_text}") # Debug
                                # else: print(f"    No linked AGD content found for {h5_text}") # Debug
                           else:
                                print(f"    Skipping already processed TSS H5: {h5_text}") # Debug

                        # Condition 2: Handle miscellaneous H5s under a needed H4 that aren't TSS/AGD
                        # Only copy these if they haven't been linked as AGD elsewhere
                        elif not h5['is_tss'] and not h5['linked_as_agd']:
                            # Check if this miscellaneous H5 has already been added
                            if h5_text not in processed_h5_texts:
                                print(f"    Adding miscellaneous needed H5: {h5_text}")
                                copy_block_to_doc(final_doc, h5['paragraph']) # Copy H5 heading
                                processed_h5_texts.add(h5_text) # Mark heading as added
                                # Copy its content unconditionally (assuming misc sections don't use <Ans#>)
                                # Or apply conditional logic if they might contain placeholders too
                                print(f"      Adding {len(h5['content'])} content block(s) for misc H5 {h5_text}")
                                for block in h5['content']:
                                    copy_block_to_doc(final_doc, block)
                            # else: print(f"    Skipping already processed misc H5: {h5_text}") # Debug
# (The loop structure continues until all H3/H4/H5 are processed)

# For highlighting the requirements that are not being satisfied in red colour.
for table in final_doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if "This requirement is not being satisfied." in paragraph.text:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 0, 0)

# Save the final document
if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)
final_doc_path = os.path.join(OUTPUT_DIR, "AAR-TSS.docx")
try:
    final_doc.save(final_doc_path)
    print(f"\nSuccessfully saved AAR-TSS document to {final_doc_path}")
except Exception as e:
    print(f"\nError saving final document: {e}")

# Create the Gaps Excel sheet
print("\nCreating Gaps Excel sheet...")
wb = Workbook()
ws = wb.active
ws.title = "Gaps"
header = ["SFR", "TSS-requirement", "Missing information"]
ws.append(header)
for cell in ws[1]:
    cell.font = Font(bold=True)

for obj in excel_data:
    ws.append([
        obj.get("SFR", ""),
        obj.get("TSS-requirement", ""),
        obj.get("Missing information", "")
    ])

excel_path = os.path.join(OUTPUT_DIR, "Gaps.xlsx")
try:
    wb.save(excel_path)
    print(f"Successfully saved Gaps sheet to {excel_path}")
except Exception as e:
    print(f"Error saving Gaps Excel sheet: {e}")

print("\nScript finished.")